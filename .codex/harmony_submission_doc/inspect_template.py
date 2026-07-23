from __future__ import annotations

import json
import sys
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def run_props(run):
    font = run.font
    rpr = run._r.rPr
    east_asia = None
    if rpr is not None and rpr.rFonts is not None:
        east_asia = rpr.rFonts.get(qn("w:eastAsia"))
    return {
        "text": run.text,
        "bold": font.bold,
        "italic": font.italic,
        "size_pt": font.size.pt if font.size else None,
        "font_name": font.name,
        "east_asia": east_asia,
        "color": str(font.color.rgb) if font.color and font.color.rgb else None,
    }


def paragraph_props(paragraph):
    fmt = paragraph.paragraph_format
    return {
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "left_indent_pt": fmt.left_indent.pt if fmt.left_indent else None,
        "right_indent_pt": fmt.right_indent.pt if fmt.right_indent else None,
        "first_line_indent_pt": fmt.first_line_indent.pt if fmt.first_line_indent else None,
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": str(fmt.line_spacing) if fmt.line_spacing is not None else None,
        "keep_with_next": fmt.keep_with_next,
        "page_break_before": fmt.page_break_before,
        "runs": [run_props(run) for run in paragraph.runs],
    }


def table_props(table):
    return {
        "style": table.style.name if table.style else None,
        "rows": [
            [
                {
                    "text": cell.text,
                    "paragraphs": [paragraph_props(p) for p in cell.paragraphs],
                }
                for cell in row.cells
            ]
            for row in table.rows
        ],
    }


def main():
    src = Path(sys.argv[1])
    out = Path(sys.argv[2])
    doc = Document(src)
    data = {
        "paragraphs": [paragraph_props(p) for p in doc.paragraphs],
        "tables": [table_props(t) for t in doc.tables],
        "sections": [
            {
                "page_width": s.page_width,
                "page_height": s.page_height,
                "top_margin": s.top_margin,
                "bottom_margin": s.bottom_margin,
                "left_margin": s.left_margin,
                "right_margin": s.right_margin,
                "header_distance": s.header_distance,
                "footer_distance": s.footer_distance,
            }
            for s in doc.sections
        ],
    }
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    with ZipFile(src) as zf:
        parts = []
        for info in zf.infolist():
            parts.append({"path": info.filename, "size": info.file_size, "crc": info.CRC})
        app_xml = ElementTree.fromstring(zf.read("docProps/app.xml"))
        pages = None
        for node in app_xml.iter():
            if node.tag.endswith("Pages"):
                pages = node.text
        document_xml = zf.read("word/document.xml")
        page_breaks = document_xml.count(b'w:type="page"')
        last_rendered = document_xml.count(b"lastRenderedPageBreak")
    out.with_name("template_package_inventory.json").write_text(
        json.dumps(parts, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    summary = [
        f"pages_cached={pages} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}",
        f"explicit_page_breaks={page_breaks} last_rendered_page_breaks={last_rendered}",
    ]
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            summary.append(f"P{i}: {paragraph.text}")
    for ti, table in enumerate(doc.tables):
        summary.append(f"\nTABLE {ti}: {len(table.rows)}x{len(table.columns)}")
        for ri, row in enumerate(table.rows):
            values = []
            for ci, cell in enumerate(row.cells):
                text = " | ".join(p.text for p in cell.paragraphs if p.text.strip())
                values.append(f"C{ci}={text}")
            summary.append(f"R{ri}: " + " || ".join(values))
    summary.append("\nBLOCK ORDER")
    paragraph_index = {id(p._p): i for i, p in enumerate(doc.paragraphs)}
    table_index = {id(t._tbl): i for i, t in enumerate(doc.tables)}
    for bi, child in enumerate(doc.element.body.iterchildren()):
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            pi = paragraph_index.get(id(child))
            summary.append(f"B{bi}: P{pi}={p.text}")
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            ti = table_index.get(id(child))
            summary.append(f"B{bi}: TABLE{ti}={len(table.rows)}x{len(table.columns)}")
    out.with_name("template_summary.txt").write_text("\n".join(summary), encoding="utf-8")


if __name__ == "__main__":
    main()
