from __future__ import annotations

import hashlib
import re
import sys
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

REFERENCE = Path(
    r"C:\Users\33517\AppData\Local\Temp\2026中国高校计算机大赛人工智能创意赛初赛（鸿蒙赛道）作品说明文档模板.docx"
)
FINAL = Path(r"E:\GuardianHub\platform\docs\GuardianHub-鸿蒙赛道初赛作品说明文档.docx")
EXPECTED_HASH = "6B858AEA863D9552AF4E7E625B6B940519DBA489E6241AEDD10F108DE0307E2F"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def assert_ok(condition, message):
    if not condition:
        raise AssertionError(message)
    print(f"PASS: {message}")


def first_table_merge_signature(path: Path):
    with ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    table = root.find(".//w:tbl", NS)
    rows = []
    for tr in table.findall("w:tr", NS):
        row = []
        for tc in tr.findall("w:tc", NS):
            tc_pr = tc.find("w:tcPr", NS)
            span = 1
            vmerge = None
            if tc_pr is not None:
                grid_span = tc_pr.find("w:gridSpan", NS)
                if grid_span is not None:
                    span = int(grid_span.get(f"{{{W_NS}}}val", "1"))
                vmerge_node = tc_pr.find("w:vMerge", NS)
                if vmerge_node is not None:
                    vmerge = vmerge_node.get(f"{{{W_NS}}}val", "continue")
            row.append((span, vmerge))
        rows.append(row)
    return rows


def package_parts(path: Path):
    with ZipFile(path) as zf:
        return set(zf.namelist())


def main():
    assert_ok(sha256(REFERENCE) == EXPECTED_HASH, "official reference remained byte-for-byte unchanged")
    assert_ok(FINAL.exists() and FINAL.stat().st_size > 100_000, "final DOCX exists and contains media")

    doc = Document(FINAL)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    combined = all_text + "\n" + table_text

    assert_ok(len(doc.sections) == 1, "final document has one section")
    ref_section = Document(REFERENCE).sections[0]
    final_section = doc.sections[0]
    geometry = (
        ref_section.page_width,
        ref_section.page_height,
        ref_section.left_margin,
        ref_section.right_margin,
        ref_section.top_margin,
        ref_section.bottom_margin,
        ref_section.header_distance,
        ref_section.footer_distance,
    )
    final_geometry = (
        final_section.page_width,
        final_section.page_height,
        final_section.left_margin,
        final_section.right_margin,
        final_section.top_margin,
        final_section.bottom_margin,
        final_section.header_distance,
        final_section.footer_distance,
    )
    assert_ok(geometry == final_geometry, "page geometry matches the official template")

    assert_ok(len(doc.tables) == 4, "team table plus three authored evidence tables are present")
    assert_ok(len(doc.inline_shapes) == 3, "three figures are embedded")
    assert_ok(first_table_merge_signature(REFERENCE) == first_table_merge_signature(FINAL),
              "team table merge topology is preserved")
    assert_ok("作品说明文档提交规范说明" not in combined, "instruction-only tail was removed")
    assert_ok("《GuardianHub 数字安全防护平台》作品原创性声明" in combined,
              "originality statement is named for the work")
    assert_ok("郑重声明：承诺本参赛队伍报名信息真实有效" in combined,
              "official originality declaration wording is retained")
    assert_ok("一、创意描述（30字以内）" in combined, "creative description section is present")
    assert_ok("二、设计稿与技术方案" in combined, "design and technical plan section is present")
    assert_ok("三、介绍文档（800字以内）" in combined, "800-character introduction section is present")
    assert_ok("四、当前实现与验证情况" in combined, "implementation evidence section is present")
    assert_ok("分享、提交和点击之前，先过一道智能安全闸门" in combined,
              "creative sentence matches the approved project positioning")
    assert_ok(len(re.sub(r"\s", "", "分享、提交和点击之前，先过一道智能安全闸门")) <= 30,
              "creative sentence stays within 30 characters")

    intro_parts = [
        p.text
        for p in doc.paragraphs
        if p.text.startswith("GuardianHub 是面向高校场景")
        or p.text.startswith("Privacy Sentinel 支持截图")
        or p.text.startswith("Link Guard 对 URL")
        or p.text.startswith("四个模块统一输出风险等级")
    ]
    intro_count = len(re.sub(r"\s", "", "".join(intro_parts)))
    assert_ok(intro_count == 783, "introduction is 783 non-whitespace characters")
    assert_ok("字数说明：正文按非空白字符计 783 字。" in combined, "introduction count is printed in the document")

    assert_ok("王**" not in combined and "南京大学" not in combined and "138****" not in combined,
              "template example identities were removed")
    assert_ok(combined.count("【待填写】") >= 20, "unknown personal fields remain visibly marked")
    assert_ok("（√）应用创新" in combined, "application innovation direction is selected")
    assert_ok("Media Library Kit" in combined and "Core File Kit" in combined
              and "Share Kit" in combined and "ArkData Preferences" in combined
              and "Network Kit" in combined,
              "five HarmonyOS native capabilities are documented")
    assert_ok("不作为当前已完成能力申报" in combined, "future plans are separated from completed capabilities")

    heading_counts = {}
    for paragraph in doc.paragraphs:
        name = paragraph.style.name if paragraph.style else ""
        if name.startswith("Heading"):
            heading_counts[name] = heading_counts.get(name, 0) + 1
    assert_ok(heading_counts.get("Heading 1", 0) == 4, "four top-level headings use real Heading 1 style")
    assert_ok(heading_counts.get("Heading 2", 0) == 4, "four second-level headings use real Heading 2 style")

    parts_ref = package_parts(REFERENCE)
    parts_final = package_parts(FINAL)
    preserved_prefixes = ("customXml/",)
    preserved_exact = {
        "word/numbering.xml",
        "word/theme/theme1.xml",
        "word/fontTable.xml",
    }
    for part in sorted(p for p in parts_ref if p.startswith(preserved_prefixes)):
        assert_ok(part in parts_final, f"preserved package part exists: {part}")
    for part in sorted(preserved_exact & parts_ref):
        assert_ok(part in parts_final, f"preserved package part exists: {part}")

    with ZipFile(FINAL) as zf:
        document_xml = zf.read("word/document.xml")
        footer_xml = zf.read("word/footer1.xml")
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
    explicit_breaks = document_xml.count(b'w:type="page"')
    assert_ok(5 <= explicit_breaks < 20, "explicit page breaks indicate a body comfortably below 20 pages")
    assert_ok(b" PAGE " in footer_xml, "footer contains a PAGE field")
    assert_ok(len(media) == 3, "package contains exactly three authored media assets")

    print("VALIDATION COMPLETE")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"FAIL: {exc}")
        sys.exit(1)
