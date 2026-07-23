from __future__ import annotations

import hashlib
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

SKILL_SCRIPTS = Path(
    r"C:\Users\33517\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


REFERENCE = Path(
    r"C:\Users\33517\AppData\Local\Temp\2026中国高校计算机大赛人工智能创意赛初赛（鸿蒙赛道）作品说明文档模板.docx"
)
EXPECTED_SHA256 = "6B858AEA863D9552AF4E7E625B6B940519DBA489E6241AEDD10F108DE0307E2F"
FINAL = Path(r"E:\GuardianHub\platform\docs\GuardianHub-鸿蒙赛道初赛作品说明文档.docx")
FIGURES = Path(r"E:\GuardianHub\.codex\harmony_submission_doc\figures")

CREATIVE = "分享、提交和点击之前，先过一道智能安全闸门"
INTRO_PARAGRAPHS = [
    "GuardianHub 是面向高校场景的数字安全防护平台，由 Privacy Sentinel 隐私哨兵升级而来。项目聚焦学生在图片分享、代码提交、链接点击和材料上传时“风险发现太晚”的问题，把检查前置到分享前、提交前、点击前、上传前四个关键节点，用一个入口提供可理解、可操作的安全决策。",
    "Privacy Sentinel 支持截图、快递单、订单页和聊天记录的本地 OCR、二维码与隐私规则检测，标注手机号、证件号、银行卡号、邮箱、地址等敏感区域，并按 high、medium、low 分级。用户可选择黑条、模糊或马赛克，预览后导出安全版本。Code Guardian 支持代码片段、单文件和项目 ZIP，检查硬编码密钥、SQL 注入、命令执行、路径穿越、弱加密、敏感日志和 XSS。项目包只在内存中读取，不解压落盘、不运行代码，自动忽略依赖与构建目录，并限制文件数、解压总量和异常压缩比；报告按文件展示语言分布、风险优先级、命中行号和修改建议。",
    "Link Guard 对 URL、短链接、二维码内容和来源场景做纯静态检查，从协议、域名、可疑参数、随机 token 与 Punycode 仿冒等维度判断是否建议打开，全程不访问目标地址。Doc Shield 可解析 PDF、DOCX、TXT、Markdown 材料，结合提交要求检查文件完整性、格式、命名和隐私风险，降低漏交、错交与信息泄露概率。",
    "四个模块统一输出风险等级、0—100 安全评分、命中证据和处置建议，检测历史保存在本地 SQLite。平台坚持“本地优先、联网可选”：核心流程可在本机或私有服务运行，只有用户主动开启时才调用可选模型增强。当前项目采用 ArkTS 构建 HarmonyOS 统一入口，已接入系统图片与文件选择器和四模块接口；React、TypeScript 与 FastAPI 组成 PC Web 验证端和检测服务，提供真实图像处理、安全上传、历史回放和自动化测试。",
]


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest().upper()


def set_run_font(run, size: float, bold: bool = False, color: str = "000000", family: str = "宋体"):
    run.font.name = family
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), family)
    rfonts.set(qn("w:hAnsi"), family)
    rfonts.set(qn("w:eastAsia"), family)
    rfonts.set(qn("w:cs"), family)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_text(
    paragraph,
    text: str,
    *,
    size: float,
    bold: bool = False,
    align=None,
    color: str = "000000",
    family: str = "宋体",
    first_indent: bool = False,
    keep_with_next: bool = False,
):
    clear_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.keep_with_next = keep_with_next
    fmt.first_line_indent = Pt(size * 2) if first_indent else Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, family=family)
    return run


def set_cover_field(paragraph, label: str, value: str, pending: bool = False):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    label_run = paragraph.add_run(label)
    set_run_font(label_run, 18, family="HarmonyOS Sans SC")
    value_run = paragraph.add_run(value)
    set_run_font(
        value_run,
        18,
        bold=not pending,
        color="C00000" if pending else "000000",
        family="HarmonyOS Sans SC",
    )
    if pending:
        value_run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def mark_pending(run):
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def cell_set_text(cell, text: str, pending: bool = False, size: float = 10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text, size=size, align=align)
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)
    if pending and paragraph.runs:
        mark_pending(paragraph.runs[0])


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def ensure_style(doc, name: str, size: float, bold: bool, color: str, before: float, after: float):
    styles = doc.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "宋体"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "宋体")
    fmt = style.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.keep_with_next = bold
    return style


def add_heading(doc, text: str, level: int):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, {1: 16, 2: 14, 3: 12}[level], bold=True, color="13233E")
    return paragraph


def add_body(doc, text: str, *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color="000000"):
    paragraph = doc.add_paragraph(style="GH 正文")
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Pt(21) if indent else Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, color=color)
    return paragraph


def set_picture_alt(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_figure(doc, path: Path, caption: str, alt: str, width_inches: float = 5.55):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width_inches))
    set_picture_alt(shape, caption, alt)
    cap = doc.add_paragraph(style="GH 图题")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, 10.5)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        header.cells[i].width = Cm(widths_cm[i])
        shade_cell(header.cells[i], "DDE8FA")
        cell_set_text(header.cells[i], text, size=10.5)
        for run in header.cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("13233E")
    for row_values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_values):
            cells[i].width = Cm(widths_cm[i])
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell_set_text(cells[i], text, size=9.5, align=align)
    exact_widths = column_widths_from_weights(widths_cm, total_width_dxa=8300)
    apply_table_geometry(table, exact_widths, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, cached, fld_end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, 9)


def set_update_fields(doc):
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def build():
    if sha256(REFERENCE) != EXPECTED_SHA256:
        raise RuntimeError("Template SHA-256 mismatch; fresh distillation is required.")

    FINAL.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(REFERENCE, FINAL)
    doc = Document(FINAL)

    # Cover slots.
    set_cover_field(doc.paragraphs[6], "参赛学校：", "【待填写】", pending=True)
    set_cover_field(doc.paragraphs[7], "团队名称：", "【待填写】", pending=True)
    set_cover_field(doc.paragraphs[8], "作品名称：", "GuardianHub 数字安全防护平台")
    set_cover_field(doc.paragraphs[9], "赛题方向：", "应用创新")
    set_cover_field(doc.paragraphs[11], "联系人（队长）：", "【待填写】", pending=True)
    set_cover_field(doc.paragraphs[12], "联系电话（队长）：", "【待填写】", pending=True)

    # Overview page.
    set_paragraph_text(doc.paragraphs[21], "作品说明文档（初赛）", size=20, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, family="HarmonyOS Sans SC")
    set_paragraph_text(doc.paragraphs[22], "本文件包含以下提交内容：", size=16, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, family="HarmonyOS Sans SC")

    # Team information table.
    table = doc.tables[0]
    cell_set_text(table.cell(0, 1), "GuardianHub 数字安全防护平台")
    cell_set_text(table.cell(1, 1), "【待填写】", pending=True)
    cell_set_text(table.cell(2, 1), "【待填写】", pending=True)
    cell_set_text(
        table.cell(3, 1),
        "（√）应用创新  （ ）Agent创新  （ ）用户体验创新  （ ）操作系统智能创新",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    member_cols = [0, 1, 2, 3, 5, 6, 8, 10, 11]
    for row_index in (6, 7, 8):
        for col in member_cols:
            value = "【待填写】"
            cell_set_text(table.cell(row_index, col), value, pending=True, size=9.5)
    teacher_cols = [0, 1, 2, 4, 7, 9]
    for col in teacher_cols:
        cell_set_text(table.cell(11, col), "【待填写】", pending=True, size=9.5)
    strengths = (
        "团队已完成 HarmonyOS 6 / ArkTS 原生客户端、四模块统一交互、FastAPI 检测服务及 PC Web 验证端；"
        "具备移动端开发、图像与文本风险识别、安全规则工程、项目 ZIP 安全扫描、接口设计和自动化测试能力。"
        "成员个人经历、具体分工及互补情况请在提交前补充。"
    )
    cell_set_text(table.cell(13, 0), strengths, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Originality statement title.
    set_paragraph_text(
        doc.paragraphs[47],
        "《GuardianHub 数字安全防护平台》作品原创性声明",
        size=20,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        family="HarmonyOS Sans SC",
    )

    # Remove instructional tail while preserving the declaration page and section properties.
    for paragraph in list(doc.paragraphs[74:]):
        remove_paragraph(paragraph)

    # Styles for authored body.
    ensure_style(doc, "Heading 1", 16, True, "13233E", 8, 6)
    ensure_style(doc, "Heading 2", 14, True, "13233E", 6, 4)
    ensure_style(doc, "Heading 3", 12, True, "13233E", 4, 3)
    ensure_style(doc, "GH 正文", 10.5, False, "000000", 0, 4)
    ensure_style(doc, "GH 图题", 10.5, False, "000000", 2, 6)

    doc.add_page_break()
    add_heading(doc, "一、创意描述（30字以内）", 1)
    creative_p = doc.add_paragraph()
    creative_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    creative_p.paragraph_format.space_before = Pt(18)
    creative_p.paragraph_format.space_after = Pt(18)
    creative_run = creative_p.add_run(CREATIVE)
    set_run_font(creative_run, 16, bold=True, color="2764E7")
    add_body(
        doc,
        "GuardianHub 将分散的校园数字风险检查统一到 HarmonyOS 原生入口，在用户执行关键动作之前完成检测、解释与处置。",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        color="617089",
    )

    add_heading(doc, "二、设计稿与技术方案", 1)
    add_heading(doc, "2.1 产品定位与交互逻辑", 2)
    add_body(
        doc,
        "产品围绕高校学生四类高频动作设计统一安全入口。用户无需理解复杂安全术语，只需在关键操作前选择图片、项目包、链接或材料，系统即返回风险等级、0—100安全评分、命中证据和下一步建议。",
    )
    add_figure(
        doc,
        FIGURES / "guardianhub_user_flow.png",
        "图 1  GuardianHub 用户交互流程",
        "四个操作前安全入口汇聚到统一安全决策中心，再输出解释、处置、继续操作与本地历史。",
    )
    add_table(
        doc,
        ["操作节点", "功能模块", "主要输入", "输出与决策"],
        [
            ["分享前", "Privacy Sentinel", "截图、订单页、快递单", "敏感区域、脱敏副本、是否建议分享"],
            ["提交前", "Code Guardian", "代码片段、单文件、项目 ZIP", "漏洞证据、文件/行号、是否建议提交"],
            ["点击前", "Link Guard", "URL、二维码、来源场景", "静态风险项、是否建议打开"],
            ["上传前", "Doc Shield", "提交要求、PDF/DOCX/TXT/MD", "完整性、格式、隐私与提交建议"],
        ],
        [2.0, 3.0, 4.1, 5.3],
    )

    doc.add_page_break()
    add_heading(doc, "2.2 技术架构与鸿蒙能力", 2)
    add_body(
        doc,
        "客户端采用 HarmonyOS 6、ArkTS 与 Stage 模型实现；检测服务由 FastAPI 提供统一接口，底层组合 OCR、二维码识别、图像处理与安全规则引擎。核心流程可运行在本机或私有服务，联网模型增强默认关闭，仅在用户主动开启并完成配置后使用。",
    )
    add_figure(
        doc,
        FIGURES / "guardianhub_architecture.png",
        "图 2  GuardianHub 技术架构与鸿蒙能力映射",
        "HarmonyOS 客户端的 Media Library Kit、Core File Kit、Share Kit、ArkData 与 Network Kit 连接四模块检测接口及底层算法和存储。",
    )
    add_table(
        doc,
        ["鸿蒙能力", "项目落地", "用户价值"],
        [
            ["Media Library Kit", "Privacy 选图、Link 二维码选图、处理图保存图库", "使用系统授权链路，减少自行管理相册权限与路径"],
            ["Core File Kit", "Code 选择 ZIP，Doc 多选材料；授权 URI 以内存方式读取", "文件边界清晰，避免无关目录访问"],
            ["Share Kit", "调用系统分享面板分享脱敏后的安全图片", "处置完成后直接回到真实分享场景"],
            ["ArkData Preferences", "保存最近检测、平均分、待复核数与模块统计", "结果留在本地，可回放、可追踪"],
            ["Network Kit", "调用四模块检测、脱敏和二维码解析接口", "统一超时、错误处理和 multipart 上传"],
        ],
        [3.0, 5.5, 5.9],
    )

    doc.add_page_break()
    add_heading(doc, "2.3 核心闭环示例", 2)
    add_body(
        doc,
        "Privacy Sentinel 是当前版本的主展示模块：系统相册选择图片后完成 OCR、二维码与隐私规则识别，用户选择黑条、模糊或马赛克生成真实安全副本，再保存到系统图库并通过 Share Kit 分享。图中所有姓名、号码和地址均为虚构测试数据。",
    )
    add_figure(
        doc,
        FIGURES / "guardianhub_privacy_example.png",
        "图 3  Privacy Sentinel 脱敏闭环示例",
        "虚构测试图片在处理前包含手机号、邮箱、学号、订单号、地址和二维码，处理后对应区域被黑条遮挡。",
    )
    add_heading(doc, "2.4 关键安全边界", 2)
    add_body(
        doc,
        "代码项目 ZIP 仅在内存中只读扫描，不执行、编译或安装依赖；系统拒绝路径穿越、符号链接、加密包和异常压缩比。Link Guard 只分析输入字符串和二维码内容，不主动访问目标网址。图片与材料上传设置大小、数量和格式限制，联网增强仅发送用户主动选择的必要内容。",
    )

    doc.add_page_break()
    add_heading(doc, "三、介绍文档（800字以内）", 1)
    for paragraph in INTRO_PARAGRAPHS:
        add_body(doc, paragraph)
    intro_text = "".join(INTRO_PARAGRAPHS)
    count = len(re.sub(r"\s", "", intro_text))
    count_p = add_body(doc, f"字数说明：正文按非空白字符计 {count} 字。", indent=False, align=WD_ALIGN_PARAGRAPH.RIGHT, color="617089")
    count_p.paragraph_format.space_before = Pt(4)

    add_heading(doc, "四、当前实现与验证情况", 1)
    add_table(
        doc,
        ["检查项", "当前状态", "说明"],
        [
            ["HarmonyOS 客户端", "已完成", "HarmonyOS 6.0.2(22)、ArkTS Stage；四模块、系统 Picker、Share Kit 与 ArkData 已在手机模拟器完成交互回归"],
            ["隐私闭环", "已完成", "真实检测、黑条/模糊/马赛克、缓存预览、保存系统图库与系统分享"],
            ["代码 ZIP 扫描", "已完成", "内存只读扫描，输出语言分布、风险文件、文件路径与行号"],
            ["链接与材料检查", "已完成", "URL 静态检查、二维码解析接口、多文件材料完整性/格式/隐私检查"],
            ["自动化验证", "已完成", "后端 66 项测试通过；PC Web 18 项测试及生产构建通过；HarmonyOS ArkTS 构建通过"],
            ["HAP", "可演示", "已生成未签名调试 HAP；正式签名包需按提交平台或真机安装要求补齐开发者 Profile"],
        ],
        [3.0, 2.6, 8.8],
    )
    add_body(
        doc,
        "真实性说明：本文“已完成”能力均可在现有工程中定位并构建；DeepSeek、Qwen VL、Scan Kit、跨设备协同和平板专项适配不作为当前已完成能力申报。",
        indent=False,
        color="617089",
    )

    for section in doc.sections:
        add_page_number(section)
    set_update_fields(doc)

    props = doc.core_properties
    props.title = "GuardianHub 鸿蒙赛道初赛作品说明文档"
    props.subject = "2026中国高校计算机大赛—人工智能创意赛鸿蒙赛道"
    props.author = "GuardianHub 参赛团队"
    props.last_modified_by = "GuardianHub 参赛团队"
    props.keywords = "GuardianHub, HarmonyOS, ArkTS, 数字安全, 高校"
    props.comments = "基于官方初赛作品说明文档模板生成"

    doc.save(FINAL)

    if sha256(REFERENCE) != EXPECTED_SHA256:
        raise RuntimeError("Reference template changed during build.")
    if len(re.sub(r"\s", "", "".join(INTRO_PARAGRAPHS))) > 800:
        raise RuntimeError("Introduction exceeds 800 characters.")
    if len(re.sub(r"\s", "", CREATIVE)) > 30:
        raise RuntimeError("Creative description exceeds 30 characters.")
    creative_count = len(re.sub(r"\s", "", CREATIVE))
    print(FINAL)
    print(f"creative_chars={creative_count}")
    print(f"intro_chars={count}")


if __name__ == "__main__":
    build()
